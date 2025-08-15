# -*- coding: utf-8 -*-
"""
API endpoints pour la gestion des conversations
"""
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from typing import List, Optional
import io
from pydantic import BaseModel, Field
from datetime import datetime
import uuid

from app.db import get_db
from app.services.conversation_service import ConversationService, MemoryService
from app.models import Conversation, Message
from app.config import settings
from openai import OpenAI

router = APIRouter()
client = OpenAI(api_key=settings.OPENAI_API_KEY)

# Modèles Pydantic pour les requêtes/réponses
class ConversationCreate(BaseModel):
    title: Optional[str] = None

class ConversationResponse(BaseModel):
    id: uuid.UUID
    title: str
    created_at: datetime
    updated_at: datetime
    is_archived: bool
    message_count: int = 0

class MessageCreate(BaseModel):
    content: str
    role: str = "user"

class MessageResponse(BaseModel):
    id: uuid.UUID
    role: str
    content: str
    created_at: datetime

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[uuid.UUID] = None
    use_memory: bool = True

class ChatResponse(BaseModel):
    message: MessageResponse
    assistant_response: MessageResponse
    conversation_id: uuid.UUID

@router.post("/conversations", response_model=ConversationResponse)
def create_conversation(
    conversation_data: ConversationCreate,
    db: Session = Depends(get_db)
):
    """Crée une nouvelle conversation"""
    service = ConversationService(db)
    conversation = service.create_conversation(conversation_data.title)
    
    return ConversationResponse(
        id=conversation.id,
        title=conversation.title,
        created_at=conversation.created_at,
        updated_at=conversation.updated_at,
        is_archived=conversation.is_archived,
        message_count=0
    )

@router.get("/conversations", response_model=List[ConversationResponse])
def get_conversations(
    limit: int = 50,
    archived: bool = False,
    db: Session = Depends(get_db)
):
    """Récupère la liste des conversations"""
    service = ConversationService(db)
    conversations = service.get_conversations(limit, archived)
    
    result = []
    for conv in conversations:
        result.append(ConversationResponse(
            id=conv.id,
            title=conv.title,
            created_at=conv.created_at,
            updated_at=conv.updated_at,
            is_archived=conv.is_archived,
            message_count=len(conv.messages)
        ))
    
    return result

@router.get("/conversations/{conversation_id}", response_model=ConversationResponse)
def get_conversation(
    conversation_id: uuid.UUID,
    db: Session = Depends(get_db)
):
    """Récupère une conversation spécifique"""
    service = ConversationService(db)
    conversation = service.get_conversation(conversation_id)
    
    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation non trouvée"
        )
    
    return ConversationResponse(
        id=conversation.id,
        title=conversation.title,
        created_at=conversation.created_at,
        updated_at=conversation.updated_at,
        is_archived=conversation.is_archived,
        message_count=len(conversation.messages)
    )

@router.get("/conversations/{conversation_id}/messages", response_model=List[MessageResponse])
def get_conversation_messages(
    conversation_id: uuid.UUID,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """Récupère les messages d'une conversation"""
    service = ConversationService(db)
    messages = service.get_conversation_messages(conversation_id, limit)
    
    return [
        MessageResponse(
            id=msg.id,
            role=msg.role,
            content=msg.content,
            created_at=msg.created_at
        )
        for msg in messages
    ]

@router.post("/conversations/{conversation_id}/messages", response_model=MessageResponse)
def add_message(
    conversation_id: uuid.UUID,
    message_data: MessageCreate,
    db: Session = Depends(get_db)
):
    """Ajoute un message à une conversation"""
    service = ConversationService(db)
    
    # Vérifier que la conversation existe
    conversation = service.get_conversation(conversation_id)
    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation non trouvée"
        )
    
    message = service.add_message(conversation_id, message_data.role, message_data.content)
    
    return MessageResponse(
        id=message.id,
        role=message.role,
        content=message.content,
        created_at=message.created_at
    )

@router.post("/chat", response_model=ChatResponse)
def chat_with_assistant(
    chat_request: ChatRequest,
    db: Session = Depends(get_db)
):
    """Chat avec l'assistant avec mémoire contextuelle"""
    service = ConversationService(db)
    memory_service = MemoryService(db)
    
    # Créer une nouvelle conversation si nécessaire
    if not chat_request.conversation_id:
        conversation = service.create_conversation()
        conversation_id = conversation.id
    else:
        conversation_id = chat_request.conversation_id
        conversation = service.get_conversation(conversation_id)
        if not conversation:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Conversation non trouvée"
            )
    
    # Ajouter le message de l'utilisateur
    user_message = service.add_message(conversation_id, "user", chat_request.message)
    
    # Récupérer le contexte de la conversation
    context = service.get_conversation_context(conversation_id)
    
    # Récupérer les mémoires pertinentes si demandé
    relevant_memories = []
    if chat_request.use_memory:
        relevant_memories = memory_service.get_relevant_memories(
            query=chat_request.message,
            limit=5
        )
    
    # Construire le prompt avec la mémoire
    system_prompt = "Tu es l'assistant de Romain. Écris en français."
    
    if relevant_memories:
        memory_context = "\n".join([
            f"Mémoire: {memory.content}" for memory in relevant_memories
        ])
        system_prompt += f"\n\nInformations pertinentes de ta mémoire:\n{memory_context}"
    
    # Préparer les messages pour OpenAI
    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(context)
    
    try:
        # Fallback local si pas de clé: simuler une réponse simple pour la démo
        if not settings.OPENAI_API_KEY:
            assistant_content = f"[LOCAL MODE] Pong: {chat_request.message}"
        else:
            # Appel à OpenAI
            response = client.chat.completions.create(
                model=settings.OPENAI_MODEL,
                messages=messages,
                temperature=0.3,
                max_tokens=1000
            )
            assistant_content = response.choices[0].message.content
        
        # Ajouter la réponse de l'assistant
        assistant_message = service.add_message(conversation_id, "assistant", assistant_content)
        
        # Analyser si de nouvelles informations doivent être mémorisées
        # (Cette logique peut être améliorée avec un modèle spécialisé)
        if any(keyword in chat_request.message.lower() for keyword in ['rappelle', 'retiens', 'important', 'note']):
            memory_service.store_memory(
                content=chat_request.message,
                context=f"Conversation du {datetime.now().strftime('%d/%m/%Y')}",
                category="user_request",
                conversation_id=conversation_id
            )
        
        return ChatResponse(
            message=MessageResponse(
                id=user_message.id,
                role=user_message.role,
                content=user_message.content,
                created_at=user_message.created_at
            ),
            assistant_response=MessageResponse(
                id=assistant_message.id,
                role=assistant_message.role,
                content=assistant_message.content,
                created_at=assistant_message.created_at
            ),
            conversation_id=conversation_id
        )
        
    except Exception as e:
        # En absence de clé, on ne devrait pas arriver ici, mais par sécurité
        if not settings.OPENAI_API_KEY:
            assistant_message = service.add_message(conversation_id, "assistant", f"[LOCAL MODE][ERROR] {type(e).__name__}")
            return ChatResponse(
                message=MessageResponse(
                    id=user_message.id,
                    role=user_message.role,
                    content=user_message.content,
                    created_at=user_message.created_at
                ),
                assistant_response=MessageResponse(
                    id=assistant_message.id,
                    role=assistant_message.role,
                    content=assistant_message.content,
                    created_at=assistant_message.created_at
                ),
                conversation_id=conversation_id
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erreur lors de l'appel à OpenAI: {str(e)}"
        )

@router.put("/conversations/{conversation_id}/title")
def update_conversation_title(
    conversation_id: uuid.UUID,
    title: str,
    db: Session = Depends(get_db)
):
    """Met à jour le titre d'une conversation"""
    service = ConversationService(db)
    success = service.update_conversation_title(conversation_id, title)
    
    if not success:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation non trouvée"
        )
    
    return {"message": "Titre mis à jour avec succès"}

@router.put("/conversations/{conversation_id}/archive")
def archive_conversation(
    conversation_id: uuid.UUID,
    db: Session = Depends(get_db)
):
    """Archive une conversation"""
    service = ConversationService(db)
    success = service.archive_conversation(conversation_id)
    
    if not success:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation non trouvée"
        )
    
    return {"message": "Conversation archivée avec succès"}

@router.delete("/conversations/{conversation_id}")
def delete_conversation(
    conversation_id: uuid.UUID,
    db: Session = Depends(get_db)
):
    """Supprime une conversation"""
    service = ConversationService(db)
    success = service.delete_conversation(conversation_id)
    
    if not success:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation non trouvée"
        )
    
    return {"message": "Conversation supprimée avec succès"}

@router.get("/conversations/{conversation_id}/export")
def export_conversation(
    conversation_id: uuid.UUID,
    format: str = "pdf",
    db: Session = Depends(get_db)
):
    """Exporte la conversation au format demandé: pdf | docx | xlsx"""
    service = ConversationService(db)
    conversation = service.get_conversation(conversation_id)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation non trouvée")

    # Collecte des messages
    rows = [
        {
            "date": m.created_at.strftime("%Y-%m-%d %H:%M"),
            "role": m.role,
            "content": m.content,
        }
        for m in conversation.messages
    ]

    fmt = (format or "pdf").lower()
    filename_base = f"conversation_{conversation.id}"

    if fmt == "pdf":
        # Construire un HTML simple et convertir en PDF via WeasyPrint
        try:
            from weasyprint import HTML, CSS
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"WeasyPrint indisponible: {e}")
        html = [
            "<html><head><meta charset='utf-8'><style>",
            "body{font-family:Arial,Helvetica,sans-serif;font-size:12pt;padding:24px;color:#0f172a}",
            "h1{font-size:16pt;margin-bottom:16px}",
            ".msg{margin:10px 0;padding:10px;border:1px solid #e2e8f0;border-radius:8px}",
            ".role{font-weight:bold;margin-bottom:6px}",
            ".date{color:#64748b;font-size:10pt}",
            "</style></head><body>",
            f"<h1>{conversation.title or 'Conversation'}</h1>",
        ]
        for r in rows:
            content_html = (r['content'] or '').replace('\n', '<br/>')
            html.append(
                f"<div class='msg'><div class='role'>{r['role'].capitalize()}</div>"
                f"<div class='date'>{r['date']}</div>"
                f"<div class='content'>{content_html}</div></div>"
            )
        html.append("</body></html>")
        html_str = "".join(html)
        pdf_bytes = HTML(string=html_str).write_pdf()
        bio = io.BytesIO(pdf_bytes)
        bio.seek(0)
        headers = {"Content-Disposition": f"attachment; filename={filename_base}.pdf"}
        return StreamingResponse(bio, media_type="application/pdf", headers=headers)

    if fmt == "docx":
        try:
            from docx import Document
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"python-docx indisponible: {e}")
        doc = Document()
        doc.add_heading(conversation.title or "Conversation", level=1)
        for r in rows:
            p = doc.add_paragraph()
            p.add_run(f"{r['role'].capitalize()} ").bold = True
            p.add_run(f"({r['date']})\n").italic = True
            for line in (r['content'] or '').splitlines():
                p.add_run(line)
                p.add_run("\n")
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        headers = {"Content-Disposition": f"attachment; filename={filename_base}.docx"}
        return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

    if fmt == "xlsx":
        try:
            import pandas as pd
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"pandas indisponible: {e}")
        df = pd.DataFrame(rows, columns=["date", "role", "content"])
        bio = io.BytesIO()
        # Use xlsxwriter if available, else default engine
        try:
            with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
                df.to_excel(writer, index=False, sheet_name="Messages")
        except Exception:
            df.to_excel(bio, index=False)
        bio.seek(0)
        headers = {"Content-Disposition": f"attachment; filename={filename_base}.xlsx"}
        return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers=headers)

    raise HTTPException(status_code=400, detail="Format non supporté. Utilisez pdf | docx | xlsx")
